"""文档导出（提示词 05）—— schemas + repository + service + 5 exporters + api 单文件版

5 种格式：PDF（WeasyPrint）/ DOCX（python-docx）/ Markdown / CSV / TXT
多文档 ZIP 打包；下载 URL HMAC 签名；过期清理。
"""

from __future__ import annotations

import os
import zipfile
from datetime import datetime, timedelta, timezone
from pathlib import Path
from typing import Literal, Protocol, cast
from urllib.parse import unquote, urlsplit
from uuid import uuid4

import structlog
from fastapi import APIRouter, Depends, Query, Request, status
from fastapi.responses import FileResponse
from pydantic import BaseModel, Field, JsonValue
from sqlalchemy import DateTime, ForeignKey, Index, Integer, String, func, select
from sqlalchemy.dialects.postgresql import ARRAY, JSONB, UUID
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import Mapped, mapped_column

from app.auth.dependencies import get_current_user, require_any_permission, require_permission
from app.common.config import settings
from app.common.database import Base, get_db
from app.common.exceptions import (
    AppException,
    ForbiddenException,
    NotFoundException,
    ValidationException,
)
from app.common.models import User
from app.common.schemas import APIResponse, ErrorCode, PaginatedData
from app.documents.models import Document, DocumentStatus
from app.documents.permissions import user_can_access_kb
from app.documents.storage import DocumentStorage
from app.exports._shared.signing import is_expired, sign_download_token, verify_download_token
from app.exports._shared.storage import delete_task_dir, root, task_file_path
from app.rag._shared.audit_helper import audit
from app.worker.queue import enqueue_export_task

logger = structlog.get_logger()
router = APIRouter(prefix="/api/v1/exports", tags=["exports"])


# ============================================================
# ORM 模型
# ============================================================
ExportFormat = Literal["pdf", "docx", "markdown", "csv", "txt"]
ExportStatus = Literal["pending", "running", "done", "failed", "expired", "cancelled"]
_CLEANUP_ELIGIBLE_STATUSES = ("done", "failed", "cancelled")


class ExportTask(Base):
    __tablename__ = "export_tasks"
    __table_args__ = (
        Index("ix_export_user_status", "user_id", "status"),
        Index("ix_export_expires", "expires_at"),
    )

    id: Mapped[str] = mapped_column(
        UUID(as_uuid=False).with_variant(String(36), "sqlite"),
        primary_key=True,
        default=lambda: str(uuid4()),
    )
    user_id: Mapped[str] = mapped_column(
        String(36), ForeignKey("users.id", ondelete="CASCADE"), nullable=False
    )
    format: Mapped[str] = mapped_column(String(16), nullable=False)
    document_ids: Mapped[list[str]] = mapped_column(
        ARRAY(UUID(as_uuid=False).with_variant(String(36), "sqlite")), nullable=False
    )
    options: Mapped[dict[str, JsonValue]] = mapped_column(JSONB, default=dict)
    status: Mapped[str] = mapped_column(String(16), default="pending")
    progress: Mapped[int] = mapped_column(Integer, default=0)
    file_path: Mapped[str | None] = mapped_column(String(1024), nullable=True)
    file_size: Mapped[int | None] = mapped_column(Integer, nullable=True)
    expires_at: Mapped[datetime] = mapped_column(DateTime(timezone=True), nullable=False)
    error_code: Mapped[str | None] = mapped_column(String(64), nullable=True)
    error_message: Mapped[str | None] = mapped_column(String(512), nullable=True)
    created_at: Mapped[datetime] = mapped_column(DateTime(timezone=True), server_default=func.now())
    finished_at: Mapped[datetime | None] = mapped_column(DateTime(timezone=True), nullable=True)


# ============================================================
# Schemas
# ============================================================
class ExportOptions(BaseModel):
    include_citations: bool = True
    include_assets: bool = True
    include_toc: bool = True
    template: Literal["default", "academic", "minimal"] = "default"
    page_size: Literal["A4", "Letter", "A3"] = "A4"
    font_size: int = 12
    language: str = "zh-CN"


class CreateExportRequest(BaseModel):
    format: ExportFormat
    document_ids: list[str] = Field(min_length=1, max_length=50)
    options: ExportOptions = Field(default_factory=ExportOptions)


class AnswerExportRequest(BaseModel):
    format: Literal["pdf", "docx", "markdown", "txt"] = "markdown"
    question: str = Field(min_length=1, max_length=2000)
    answer: str = Field(min_length=1, max_length=100_000)
    citations: list[dict[str, JsonValue]] = Field(default_factory=list, max_length=50)


class ExportTaskResponse(BaseModel):
    id: str
    user_id: str
    format: ExportFormat
    document_ids: list[str]
    options: ExportOptions
    status: ExportStatus
    progress: int
    source_type: Literal["document", "answer"]
    filename: str | None = None
    file_size: int | None = None
    download_url: str | None = None
    expires_at: datetime
    error_code: str | None = None
    error_message: str | None = None
    created_at: datetime | None = None
    finished_at: datetime | None = None
    model_config = {"from_attributes": True}


# ============================================================
# Exporter 抽象与实现（提示词 05 §4.2）
# ============================================================
class ExportContent:
    def __init__(
        self,
        document_id: str,
        title: str,
        markdown: str,
        assets: list[Path] | None = None,
        citations: list[dict[str, JsonValue]] | None = None,
        metadata: dict[str, JsonValue] | None = None,
    ) -> None:
        self.document_id = document_id
        self.title = title
        self.markdown = markdown
        self.assets = assets or []
        self.citations = citations or []
        self.metadata = metadata or {}


class Exporter(Protocol):
    format_name: str

    async def export(
        self, content: ExportContent, output_path: str, options: ExportOptions
    ) -> str: ...


class MarkdownExporter:
    format_name = "markdown"

    async def export(self, content: ExportContent, output_path: str, options: ExportOptions) -> str:
        body = f"# {content.title}\n\n{content.markdown}\n"
        if options.include_citations and content.citations:
            body += "\n## 引用\n\n" + "\n".join(
                f"- doc={c.get('doc_id')} chunk={c.get('chunk_id')} score={c.get('score')}"
                for c in content.citations
            )
        Path(output_path).write_text(body, encoding="utf-8")
        return output_path


class TxtExporter:
    format_name = "txt"

    async def export(self, content: ExportContent, output_path: str, options: ExportOptions) -> str:
        from bs4 import BeautifulSoup
        from markdown_it import MarkdownIt

        html = MarkdownIt().render(content.markdown)
        text = BeautifulSoup(html, "lxml").get_text("\n", strip=True)
        Path(output_path).write_text(f"{content.title}\n\n{text}", encoding="utf-8")
        return output_path


class CsvExporter:
    format_name = "csv"

    async def export(self, content: ExportContent, output_path: str, options: ExportOptions) -> str:
        import csv

        with open(output_path, "w", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)
            writer.writerow(["section", "content"])
            writer.writerow(["title", _safe_csv_cell(content.title)])
            for line in content.markdown.split("\n"):
                writer.writerow(["body", _safe_csv_cell(line)])
        return output_path


class DocxExporter:
    format_name = "docx"

    async def export(self, content: ExportContent, output_path: str, options: ExportOptions) -> str:
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("python-docx 未安装") from exc
        doc = Document()
        doc.add_heading(content.title, level=1)
        for line in content.markdown.split("\n"):
            if line.startswith("# "):
                doc.add_heading(line[2:], level=2)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=3)
            elif line.strip():
                doc.add_paragraph(line)
        if options.include_citations and content.citations:
            doc.add_heading("引用", level=2)
            for c in content.citations:
                doc.add_paragraph(
                    f"doc={c.get('doc_id')} chunk={c.get('chunk_id')} score={c.get('score')}",
                    style="List Bullet",
                )
        doc.save(output_path)
        return output_path


class PdfExporter:
    format_name = "pdf"

    async def export(self, content: ExportContent, output_path: str, options: ExportOptions) -> str:
        try:
            from weasyprint import HTML  # type: ignore[import-untyped,unused-ignore]
        except ImportError as exc:
            raise RuntimeError("weasyprint 未安装") from exc
        body_html = _markdown_to_html(content.markdown)
        html = f"""<!DOCTYPE html><html><head><meta charset="utf-8"><style>
        @page {{ size: {options.page_size}; margin: 2cm; }}
        body {{
            font-family: 'Noto Sans CJK SC', sans-serif;
            font-size: {options.font_size}pt;
            line-height: 1.6;
        }}
        h1 {{ font-size: {options.font_size + 6}pt; }}
        h2 {{ font-size: {options.font_size + 4}pt; }}
        </style></head><body>
        <h1>{_escape(content.title)}</h1>
        {body_html}
        </body></html>"""
        HTML(
            string=html,
            base_url=str(root()),
            url_fetcher=_local_asset_url_fetcher,
        ).write_pdf(output_path)
        return output_path


def _safe_csv_cell(value: str) -> str:
    """阻止电子表格把文档正文解释为公式。"""
    if value.startswith(("=", "+", "-", "@", "\t", "\r")):
        return f"'{value}"
    return value


def _local_asset_url_fetcher(
    url: str,
    timeout: int = 10,
    ssl_context: object | None = None,
) -> dict[str, object]:
    """PDF 渲染只可读取导出存储根目录内的本地文件。"""
    parsed = urlsplit(url)
    if parsed.scheme != "file":
        raise ValueError("PDF 导出禁止访问网络资源")
    candidate = Path(unquote(parsed.path)).resolve()
    storage_root = root().resolve()
    if not candidate.is_relative_to(storage_root) or not candidate.is_file():
        raise ValueError("PDF 导出资源不在受控存储目录")

    from weasyprint import default_url_fetcher

    result = default_url_fetcher(url, timeout=timeout, ssl_context=ssl_context)
    return cast(dict[str, object], result)


def _markdown_to_html(md_text: str) -> str:
    try:
        from markdown_it import MarkdownIt
    except ImportError:
        return f"<pre>{_escape(md_text)}</pre>"
    rendered: object = MarkdownIt().render(md_text)
    if not isinstance(rendered, str):
        raise TypeError("Markdown renderer returned a non-string value")
    return rendered


def _escape(s: str) -> str:
    return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


EXPORTERS: dict[str, Exporter] = {
    "markdown": MarkdownExporter(),
    "txt": TxtExporter(),
    "csv": CsvExporter(),
    "docx": DocxExporter(),
    "pdf": PdfExporter(),
}


# ============================================================
# ZIP 打包
# ============================================================
def zip_documents(*, task_id: str, files: list[str]) -> str:
    """打包多文档为单一 ZIP；返回 zip 路径。"""
    out_path = task_file_path(task_id, f"{task_id}.zip")
    with zipfile.ZipFile(out_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for fp in files:
            arcname = os.path.basename(fp)
            zf.write(fp, arcname=arcname)
    return out_path


# ============================================================
# 文档内容获取
# ============================================================
async def _fetch_document_content(
    db: AsyncSession,
    document_id: str,
    *,
    user: User | None = None,
) -> ExportContent:
    """从已处理文档的受控存储读取真实 Markdown。"""
    document = await db.get(Document, document_id)
    if document is None:
        raise NotFoundException(message="导出文档不存在")
    if document.status != DocumentStatus.READY.value:
        raise ValidationException(message=f"文档 {document_id} 尚未处理完成")
    if user is not None and not await user_can_access_kb(db, user, document.knowledge_base_id):
        raise ForbiddenException(message="导出权限已被撤销")
    storage = DocumentStorage()
    return ExportContent(
        document_id=document_id,
        title=document.title,
        markdown=storage.read_markdown(document_id),
        metadata=storage.read_manifest(document_id),
    )


# ============================================================
# 任务执行
# ============================================================
async def _run_export(db: AsyncSession, task_id: str) -> None:
    """执行导出任务（arq 异步 / 测试同步调用都复用此函数）。"""
    task = await db.get(ExportTask, task_id)
    if task is None:
        return
    if task.status in {"done", "expired", "cancelled"}:
        return
    delete_task_dir(task_id)
    task.status = "running"
    task.progress = 0
    await db.commit()

    files: list[str] = []
    try:
        for i, doc_id in enumerate(task.document_ids):
            user_result = await db.execute(
                select(User)
                .where(User.id == task.user_id)
                .execution_options(populate_existing=True)
            )
            task_user = user_result.scalar_one_or_none()
            if task_user is None or task_user.status != "active":
                raise ForbiddenException(message="导出用户已失效")
            content = await _fetch_document_content(db, doc_id, user=task_user)
            exporter = EXPORTERS.get(task.format)
            if exporter is None:
                raise ValidationException(message=f"unsupported format: {task.format}")
            opts = ExportOptions.model_validate(task.options or {})
            ext = task.format if task.format != "markdown" else "md"
            filename = f"{i + 1:02d}_{doc_id[:8]}.{ext}"
            output_path = task_file_path(task_id, filename)
            await exporter.export(content, output_path, opts)
            files.append(output_path)
            task.progress = int((i + 1) / len(task.document_ids) * 100)
            await db.commit()

        # 单文档直接返回；多文档打 ZIP
        if len(files) == 1:
            final = files[0]
        else:
            final = zip_documents(task_id=task_id, files=files)

        task.file_path = final
        task.file_size = os.path.getsize(final)
        task.status = "done"
        task.finished_at = datetime.now(timezone.utc)
    except Exception as exc:
        logger.exception("export_failed", task_id=task_id, error_type=type(exc).__name__)
        task.status = "failed"
        task.error_code = "export_failed"
        task.error_message = "导出任务执行失败"
        task.finished_at = datetime.now(timezone.utc)
    await db.commit()


# ============================================================
# 过期清理（提示词 05 §4.7）
# ============================================================
async def cleanup_expired(db: AsyncSession) -> int:
    """锁定并清理已过期的终态任务，避免与仍在执行的任务竞争。"""
    now = datetime.now(timezone.utc)
    res = await db.execute(
        select(ExportTask)
        .where(
            ExportTask.expires_at < now,
            ExportTask.status.in_(_CLEANUP_ELIGIBLE_STATUSES),
        )
        .with_for_update(skip_locked=True)
    )
    tasks = res.scalars().all()
    count = 0
    for task in tasks:
        if task.status not in _CLEANUP_ELIGIBLE_STATUSES:
            continue
        try:
            delete_task_dir(task.id)
        except (OSError, ValueError) as exc:
            logger.warning(
                "export_cleanup_failed",
                task_id=task.id,
                error_type=type(exc).__name__,
            )
            continue
        task.status = "expired"
        count += 1
    await db.commit()
    return count


def _resolve_download_path(task: ExportTask) -> Path:
    """只返回当前导出任务目录内实际存在的普通文件。"""
    storage_root = root().resolve()
    task_root = (storage_root / task.id).resolve()
    if task_root.parent != storage_root or not task.file_path:
        raise NotFoundException()
    try:
        candidate = Path(task.file_path).resolve(strict=True)
    except OSError as exc:
        raise NotFoundException() from exc
    if not candidate.is_relative_to(task_root) or not candidate.is_file():
        raise NotFoundException()
    return candidate


# ============================================================
# 路由
# ============================================================
def _task_to_response(task: ExportTask) -> ExportTaskResponse:
    download_url = None
    if task.status == "done" and task.file_path:
        expires = int(task.expires_at.timestamp())
        token = sign_download_token(
            export_id=task.id, user_id=task.user_id, expires_at_unix=expires
        )
        download_url = f"/api/v1/exports/{task.id}/download?token={token}&expires={expires}"
    return ExportTaskResponse.model_validate(
        {
            "id": task.id,
            "user_id": task.user_id,
            "format": task.format,
            "document_ids": task.document_ids,
            "options": task.options or {},
            "status": task.status,
            "progress": task.progress,
            "source_type": "answer" if not task.document_ids else "document",
            "filename": Path(task.file_path).name if task.file_path else None,
            "file_size": task.file_size,
            "download_url": download_url,
            "expires_at": task.expires_at,
            "error_code": task.error_code,
            "error_message": task.error_message,
            "created_at": task.created_at,
            "finished_at": task.finished_at,
        }
    )


@router.post("/answer")
async def export_answer_endpoint(
    request: Request,
    payload: AnswerExportRequest,
    user: User = Depends(get_current_user),
    _perm: None = Depends(require_any_permission("export:write", "export.create")),
    db: AsyncSession = Depends(get_db),
) -> FileResponse:
    """导出当前 RAG 问答并保留可再次下载的任务记录。"""
    options = ExportOptions(
        include_citations=bool(payload.citations),
        include_assets=False,
        include_toc=False,
    )
    task = ExportTask(
        user_id=user.id,
        format=payload.format,
        document_ids=[],
        options=options.model_dump(),
        status="running",
        progress=0,
        expires_at=datetime.now(timezone.utc)
        + timedelta(hours=settings.export_default_ttl_hours),
    )
    db.add(task)
    await db.commit()

    extension = "md" if payload.format == "markdown" else payload.format
    filename = f"RAG-answer-{datetime.now(timezone.utc).strftime('%Y%m%d-%H%M%S')}.{extension}"
    output_path = task_file_path(task.id, filename)
    content = ExportContent(
        document_id="rag-answer",
        title="RAG 问答结果",
        markdown=(
            f"## 问题\n\n{payload.question.strip()}\n\n"
            f"## 答案\n\n{payload.answer.strip()}"
        ),
        citations=payload.citations,
    )
    try:
        await EXPORTERS[payload.format].export(
            content,
            str(output_path),
            options,
        )
        task.file_path = str(output_path)
        task.file_size = os.path.getsize(output_path)
        task.status = "done"
        task.progress = 100
        task.finished_at = datetime.now(timezone.utc)
        await db.commit()
    except Exception as exc:
        delete_task_dir(task.id)
        task.status = "failed"
        task.error_code = "export_failed"
        task.error_message = "问答结果导出失败"
        task.finished_at = datetime.now(timezone.utc)
        await db.commit()
        logger.warning(
            "answer_export_failed",
            task_id=task.id,
            format=payload.format,
            error_type=type(exc).__name__,
        )
        raise AppException(
            code=ErrorCode.EXPORT_FAILED,
            message="问答结果导出失败，请稍后重试",
            status_code=500,
        ) from exc
    await audit(
        db,
        action="answer_export",
        user_id=user.id,
        resource_type="rag_answer",
        resource_id=task.id,
        request=request,
    )
    await db.commit()
    return FileResponse(
        path=output_path,
        filename=filename,
        media_type={
            "pdf": "application/pdf",
            "markdown": "text/markdown; charset=utf-8",
            "txt": "text/plain; charset=utf-8",
            "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        }[payload.format],
        headers={"X-Export-Id": task.id},
    )


@router.post("", status_code=status.HTTP_202_ACCEPTED)
async def create_export_endpoint(
    request: Request,
    payload: CreateExportRequest,
    user: User = Depends(get_current_user),
    _perm: None = Depends(require_any_permission("export:write", "export.create")),
    db: AsyncSession = Depends(get_db),
) -> APIResponse[ExportTaskResponse]:
    if len(set(payload.document_ids)) != len(payload.document_ids):
        raise ValidationException(message="document_ids 不能包含重复项")
    for document_id in payload.document_ids:
        document = await db.get(Document, document_id)
        if document is None:
            raise NotFoundException(message="导出文档不存在")
        if not await user_can_access_kb(db, user, document.knowledge_base_id):
            raise ForbiddenException(message="无权导出所选文档")
        if document.status != DocumentStatus.READY.value:
            raise ValidationException(message=f"文档 {document_id} 尚未处理完成")

    task = ExportTask(
        user_id=user.id,
        format=payload.format,
        document_ids=payload.document_ids,
        options=payload.options.model_dump(),
        status="pending",
        progress=0,
        expires_at=datetime.now(timezone.utc) + timedelta(hours=settings.export_default_ttl_hours),
    )
    db.add(task)
    await db.commit()
    await audit(
        db,
        action="export_create",
        user_id=user.id,
        resource_type="export_task",
        resource_id=task.id,
        request=request,
    )
    await db.commit()

    if settings.worker_inline:
        await _run_export(db, task.id)
    else:
        try:
            await enqueue_export_task(task.id, settings.redis_url)
        except Exception as exc:
            task.status = "failed"
            task.error_code = "queue_unavailable"
            task.error_message = "导出任务未能加入处理队列"
            task.finished_at = datetime.now(timezone.utc)
            await db.commit()
            logger.warning(
                "export_enqueue_failed",
                task_id=task.id,
                error_type=type(exc).__name__,
            )
            raise AppException(
                code=ErrorCode.EXPORT_FAILED,
                message="任务队列暂不可用，请稍后重试",
                status_code=503,
            ) from exc
    return APIResponse[ExportTaskResponse](data=_task_to_response(task))


@router.get("")
async def list_exports_endpoint(
    request: Request,
    page: int = Query(1, ge=1),
    page_size: int = Query(20, ge=1, le=100),
    status_filter: str | None = Query(None, alias="status"),
    user: User = Depends(get_current_user),
    _perm: None = Depends(require_any_permission("export:read", "export.view")),
    db: AsyncSession = Depends(get_db),
) -> APIResponse[PaginatedData[ExportTaskResponse]]:
    q = select(ExportTask).where(ExportTask.user_id == user.id)
    if status_filter:
        q = q.where(ExportTask.status == status_filter)
    total = (await db.execute(select(func.count()).select_from(q.subquery()))).scalar() or 0
    offset = (page - 1) * page_size
    res = await db.execute(q.order_by(ExportTask.created_at.desc()).offset(offset).limit(page_size))
    items = [_task_to_response(t) for t in res.scalars().all()]
    return APIResponse[PaginatedData[ExportTaskResponse]](
        data=PaginatedData[ExportTaskResponse](
            items=items,
            page=page,
            page_size=page_size,
            total=total,
        )
    )


@router.get("/{export_id}")
async def get_export_endpoint(
    export_id: str,
    request: Request,
    user: User = Depends(get_current_user),
    _perm: None = Depends(require_any_permission("export:read", "export.view")),
    db: AsyncSession = Depends(get_db),
) -> APIResponse[ExportTaskResponse]:
    task = await db.get(ExportTask, export_id)
    if task is None:
        raise NotFoundException()
    if task.user_id != user.id:
        raise ForbiddenException()
    return APIResponse[ExportTaskResponse](data=_task_to_response(task))


@router.delete("/{export_id}")
async def delete_export_endpoint(
    export_id: str,
    request: Request,
    user: User = Depends(get_current_user),
    _perm: None = Depends(require_permission("export:delete")),
    db: AsyncSession = Depends(get_db),
) -> APIResponse[None]:
    task = await db.get(ExportTask, export_id)
    if task is None:
        raise NotFoundException()
    if task.user_id != user.id:
        raise ForbiddenException()
    delete_task_dir(task.id)
    await db.delete(task)
    await db.commit()
    await audit(db, action="export_delete", user_id=user.id, resource_id=export_id, request=request)
    await db.commit()
    return APIResponse[None](data=None)


@router.get("/{export_id}/download")
async def download_export_endpoint(
    export_id: str,
    request: Request,
    token: str = Query(...),
    expires: int = Query(...),
    user: User = Depends(get_current_user),
    _perm: None = Depends(require_permission("export:download")),
    db: AsyncSession = Depends(get_db),
) -> FileResponse:
    task = await db.get(ExportTask, export_id)
    if task is None:
        raise NotFoundException()
    if task.user_id != user.id:
        raise ForbiddenException(message="无权下载此文件")
    for document_id in task.document_ids:
        await _fetch_document_content(db, document_id, user=user)
    if is_expired(expires):
        raise NotFoundException(code=16001, message="下载链接已过期")
    if not verify_download_token(
        export_id=task.id, user_id=task.user_id, expires_at_unix=expires, token=token
    ):
        raise ForbiddenException(message="签名无效")
    download_path = _resolve_download_path(task)
    await audit(
        db, action="export_download", user_id=user.id, resource_id=export_id, request=request
    )
    await db.commit()
    return FileResponse(
        download_path,
        filename=download_path.name,
        media_type="application/octet-stream",
    )


@router.post("/cleanup")
async def cleanup_endpoint(
    request: Request,
    user: User = Depends(get_current_user),
    _perm: None = Depends(require_permission("admin.export.cleanup")),
    db: AsyncSession = Depends(get_db),
) -> APIResponse[dict[str, int]]:
    """管理员触发过期清理；Worker 也会定时调用同一清理函数。"""
    count = await cleanup_expired(db)
    await audit(
        db, action="export_cleanup", user_id=user.id, detail=f"cleaned {count}", request=request
    )
    await db.commit()
    return APIResponse[dict[str, int]](data={"cleaned": count})
